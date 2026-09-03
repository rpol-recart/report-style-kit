#!/usr/bin/env python3
# -*- coding: utf-8 -*-
r"""
lint_text.py — Task 6 (Step 6.1) комплекта report-style-kit. Стадия 1 линта.

Проверяет тексты манифеста (content-manifest.json) или собранного .docx
против базы анти-сленга и (опционально) стилевого профиля.

CLI:
    python3 lint_text.py --manifest m.json [--profile p.json] [--glossary g.md] \
                         [--blacklist b.txt] [--out report.txt]
    python3 lint_text.py --docx d.docx   [--profile p.json] [--glossary g.md] \
                         [--blacklist b.txt] [--out report.txt]

Что проверяется:
  FAIL (по 1 сообщению на категорию, цитата ≤60 символов + локатор):
    все категории blacklist-файла (формат: категория<TAB>regex, # комментарии):
    placeholders, first_person, agent_phrases, marketing, office_anglicisms,
    emoji, markdown_residue, exclamations.
    Паттерны применяются к каждому абзацу/ячейке отдельно, поэтому ^ в
    markdown_residue привязан к началу текста абзаца/ячейки.
  WARN (не влияют на exit-код):
    glossary      — латинские токены, отсутствующие в glossary_terms профиля
                    (и/или в --glossary md); список уникальных + локатор
                    первого вхождения;
    para_length   — абзац длиннее 5 предложений (сплит по [.!?…] + пробел);
    typography    — против profile.microtypography: десятичная запятая при
                    decimal_separator=point; «\d%» при percent_spacing=spaced
                    и «\d %» при glued (при mixed/tight не проверяется);
                    прямая кавычка " при quotes=guillemets.
    Без --profile проверки glossary/typography пропускаются.

Отчёт: человекочитаемый (рус.) в stdout + файл lint_report.txt рядом с
целевым файлом (или путь из --out).

Коды выхода: 0 — PASS (допустимы предупреждения); 1 — FAIL; 2 — ошибки
использования/IO.

Зависимости: python-docx (для режима --docx). Python 3.12.
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

import docx

# Тот же токенизатор латиницы, что в extract_profile.py (Step 4.1).
TOKEN_RE = re.compile(r"[A-Za-z_][A-Za-z0-9_\-\.]{1,}")
SENT_SPLIT_RE = re.compile(r"(?<=[.!?…])\s+")
MAX_QUOTE = 60


# --------------------------------------------------------------------------- inputs

def display_path(path: Path) -> str:
    """Путь для отчёта: относительный от cwd, когда файл внутри него
    (автономность контура — без абсолютных путей хоста в артефактах)."""
    try:
        return str(path.resolve().relative_to(Path.cwd().resolve()))
    except ValueError:
        return str(path)


def load_blacklist(path: Path) -> list:
    """[(категория, compiled regex)] из файла форматом категория<TAB>regex."""
    cats = []
    for raw in path.read_text(encoding="utf-8").splitlines():
        if not raw.strip() or raw.lstrip().startswith("#") or "\t" not in raw:
            continue
        cat, rx = raw.split("\t", 1)
        try:
            cats.append((cat.strip(), re.compile(rx.strip(), re.IGNORECASE)))
        except re.error as exc:
            print(f"[warn] категория {cat.strip()!r}: некорректный regex ({exc}) — "
                  f"пропущена", file=sys.stderr)
    return cats


def iter_manifest_texts(m: dict):
    """(локатор, текст) по всем text/label/items/caption/explanation/header+rows."""
    for i, sec in enumerate(m.get("sections") or []):
        for j, b in enumerate(sec.get("blocks") or []):
            base = f"sections[{i}].blocks[{j}]"
            t = b.get("type")
            if t in ("thesis", "formula", "verdict"):
                yield f"{base}.text", b.get("text", "")
            elif t == "lead_in":
                yield f"{base}.label", b.get("label", "")
                yield f"{base}.text", b.get("text", "")
            elif t == "bullets":
                for k, it in enumerate(b.get("items") or []):
                    yield f"{base}.items[{k}]", it
            elif t == "table":
                for c, h in enumerate(b.get("header") or []):
                    yield f"{base}.header[{c}]", h
                for r, row in enumerate(b.get("rows") or []):
                    for c, cell in enumerate(row):
                        yield f"{base}.rows[{r}][{c}]", cell
            elif t == "figure":
                yield f"{base}.caption", b.get("caption", "")
                if b.get("explanation"):
                    yield f"{base}.explanation", b["explanation"]


def iter_docx_texts(d):
    """(локатор, текст) по всем непустым абзацам и ячейкам таблиц docx."""
    for n, p in enumerate(d.paragraphs):
        if p.text.strip():
            yield f"para#{n}", p.text
    for ti, t in enumerate(d.tables):
        for r, row in enumerate(t.rows):
            for c, cell in enumerate(row.cells):
                if cell.text.strip():
                    yield f"table#{ti}/r{r}/c{c}", cell.text


def parse_glossary_md(path: Path) -> set:
    """Термины из черновика глоссария (md-таблица «Термин | Частота | …»)."""
    terms = set()
    for line in path.read_text(encoding="utf-8").splitlines():
        s = line.strip()
        if not s.startswith("|"):
            continue
        cells = [c.strip() for c in s.strip("|").split("|")]
        if not cells or not cells[0] or set(cells[0]) <= {"-", " ", ":"}:
            continue  # шапка/разделитель
        if cells[0].lower() == "термин":
            continue
        terms.add(cells[0].lower())
    return terms


# --------------------------------------------------------------------------- checks

def quote_of(text: str, start: int, end: int) -> str:
    """Цитата вокруг матча, ≤60 символов, переносы строк заменены."""
    lo = max(0, start - 15)
    hi = min(len(text), end + 45)
    q = text[lo:hi].replace("\n", " ")
    if lo > 0:
        q = "…" + q
    if hi < len(text):
        q += "…"
    return q[:MAX_QUOTE]


def scan_categories(categories: list, texts: list) -> list:
    """FAIL-записи: по 1 сообщению на категорию (первый матч)."""
    fails = []
    for cat, rx in categories:
        hit = None
        n_texts = 0
        for loc, text in texts:
            m = rx.search(text)
            if m:
                n_texts += 1
                if hit is None:
                    hit = (loc, text, m)
        if hit:
            loc, text, m = hit
            fails.append({
                "category": cat,
                "locator": loc,
                "quote": quote_of(text, m.start(), m.end()),
                "n_texts": n_texts,
            })
    return fails


def scan_glossary(texts: list, allowed: set) -> list:
    """Уникальные латинские токены вне глоссария + локатор первого вхождения."""
    unknown: dict = {}
    for loc, text in texts:
        for m in TOKEN_RE.finditer(text):
            tok = m.group(0).rstrip(".").lower()
            if len(tok) >= 2 and tok not in allowed:
                unknown.setdefault(tok, loc)
    return [(tok, loc) for tok, loc in sorted(unknown.items())]


def scan_para_length(texts: list, limit: int = 5) -> list:
    warns = []
    for loc, text in texts:
        n = len([s for s in SENT_SPLIT_RE.split(text) if s.strip()])
        if n > limit:
            warns.append((loc, n))
    return warns


def scan_typography(texts: list, micro: dict) -> list:
    warns = []

    def first_match(rx: re.Pattern, label: str):
        for loc, text in texts:
            m = rx.search(text)
            if m:
                warns.append(f"{label} — {loc}: «{quote_of(text, m.start(), m.end())}»")
                return

    if micro.get("decimal_separator") == "point":
        first_match(re.compile(r"\d,\d"),
                    "[typography] десятичная запятая при decimal_separator=point")
    ps = micro.get("percent_spacing")
    if ps == "spaced":
        first_match(re.compile(r"\d%"),
                    "[typography] процент без пробела («\\d%») при percent_spacing=spaced")
    elif ps == "glued":
        first_match(re.compile(r"\d\s+%"),
                    "[typography] процент с пробелом («\\d %») при percent_spacing=glued")
    if micro.get("quotes") == "guillemets":
        first_match(re.compile("\u0022"),
                    "[typography] прямая кавычка (\") при quotes=guillemets")
    return warns


# --------------------------------------------------------------------------- report

def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(
        description="Линт текста манифеста или docx против базы анти-сленга и профиля.")
    mode = ap.add_mutually_exclusive_group(required=True)
    mode.add_argument("--manifest", help="content-manifest.json")
    mode.add_argument("--docx", help="собранный .docx")
    ap.add_argument("--profile", help="style-profile.json (включает glossary/typography)")
    ap.add_argument("--glossary", help="glossary.draft.md — доп. разрешённые термины")
    ap.add_argument("--blacklist", help="файл категория<TAB>regex "
                                         "(по умолчанию base/slang-blacklist.txt)")
    ap.add_argument("--out", help="путь отчёта (по умолчанию lint_report.txt рядом "
                                  "с целевым файлом)")
    args = ap.parse_args(argv)

    kit_root = Path(__file__).resolve().parent.parent
    target = Path(args.manifest or args.docx)
    if not target.is_file():
        print(f"[ошибка] целевой файл не найден: {target}", file=sys.stderr)
        return 2

    blacklist_path = Path(args.blacklist) if args.blacklist \
        else kit_root / "base" / "slang-blacklist.txt"
    if not blacklist_path.is_file():
        print(f"[ошибка] blacklist не найден: {blacklist_path}", file=sys.stderr)
        return 2
    categories = load_blacklist(blacklist_path)
    if not categories:
        print(f"[ошибка] в blacklist нет активных категорий: {blacklist_path}",
              file=sys.stderr)
        return 2

    profile = None
    if args.profile:
        if not Path(args.profile).is_file():
            print(f"[ошибка] профиль не найден: {args.profile}", file=sys.stderr)
            return 2
        try:
            profile = json.loads(Path(args.profile).read_text(encoding="utf-8"))
        except json.JSONDecodeError as exc:
            print(f"[ошибка] профиль — некорректный JSON: {exc}", file=sys.stderr)
            return 2

    glossary_terms: set = set()
    if profile:
        glossary_terms |= {g.get("term", "").lower()
                           for g in profile.get("glossary_terms") or [] if g.get("term")}
    if args.glossary:
        if not Path(args.glossary).is_file():
            print(f"[ошибка] глоссарий не найден: {args.glossary}", file=sys.stderr)
            return 2
        glossary_terms |= parse_glossary_md(Path(args.glossary))

    # --- извлечение текстов ----------------------------------------------------
    if args.manifest:
        try:
            m = json.loads(target.read_text(encoding="utf-8"))
        except json.JSONDecodeError as exc:
            print(f"[ошибка] манифест — некорректный JSON: {exc}", file=sys.stderr)
            return 2
        texts = list(iter_manifest_texts(m))
    else:
        d = docx.Document(str(target))
        texts = list(iter_docx_texts(d))

    if not texts:
        print("[ошибка] не найдено ни одного текста для проверки", file=sys.stderr)
        return 2

    # --- FAIL: категории блэклиста ---------------------------------------------
    fails = scan_categories(categories, texts)

    # --- WARN -------------------------------------------------------------------
    warns: list = []
    if profile or args.glossary:
        for tok, loc in scan_glossary(texts, glossary_terms):
            warns.append(f"[glossary] латинский токен «{tok}» отсутствует в глоссарии — {loc}")
    for loc, n in scan_para_length(texts):
        warns.append(f"[para_length] абзац из {n} предложений (>5) — {loc}")
    if profile:
        for w in scan_typography(texts, profile.get("microtypography") or {}):
            warns.append(w)

    # --- отчёт -------------------------------------------------------------------
    lines = [
        "=" * 72,
        f"lint_text: {display_path(target)}",
        f"Режим: {'manifest' if args.manifest else 'docx'} | "
        f"Blacklist: {display_path(blacklist_path)} | "
        f"Профиль: {args.profile or '—'} | Глоссарий: {args.glossary or '—'}",
        f"Проверено текстов: {len(texts)}",
        "=" * 72,
        f"НАРУШЕНИЯ (FAIL): {len(fails)}",
    ]
    for n, f in enumerate(fails, start=1):
        lines += [
            f"  {n}. [{f['category']}] нарушение в {f['n_texts']} текст(ах)",
            f"     где: {f['locator']}",
            f"     цитата: «{f['quote']}»",
        ]
    lines.append(f"ПРЕДУПРЕЖДЕНИЯ (warn): {len(warns)}")
    for n, w in enumerate(warns, start=1):
        lines.append(f"  {n}. {w}")

    report_path = Path(args.out) if args.out else target.resolve().parent / "lint_report.txt"
    report_path.parent.mkdir(parents=True, exist_ok=True)
    if fails:
        cats = ", ".join(f["category"] for f in fails)
        lines.append(f"ИТОГ: FAIL: {len(fails)} нарушений (категории: {cats})")
    else:
        lines.append(f"ИТОГ: PASS с {len(warns)} предупреждениями")
    report_path.write_text("\n".join(lines) + "\n", encoding="utf-8")

    print("\n".join(lines))
    return 1 if fails else 0


if __name__ == "__main__":
    sys.exit(main())
