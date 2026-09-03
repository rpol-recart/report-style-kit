#!/usr/bin/env python3
# -*- coding: utf-8 -*-
r"""
lint_visual.py — Task 7 (Step 7.1) комплекта report-style-kit. Стадия 2 линта.

Проверяет визуальное соответствие собранного .docx стилевому профилю
(style-profile.json): стили абзацев, титульный блок, последовательность
заголовков, таблицы, шрифты и параметры страницы.

CLI:
    python3 lint_visual.py --docx d.docx --profile p.json [--out report.txt]

Проверки (каждая даёт 0 или более сообщений; в сообщении — название проверки,
локатор para#N / table#T и детали):
  1. Стили абзацев   — каждый непустой абзац использует стиль из множества
                       {Normal, Heading 1, Heading 2, Heading 3, List Bullet},
                       иначе «недопустимый стиль 'X'». Пустые абзацы
                       пропускаются.
  2. Титульный блок  — первые непустые абзацы до первого Heading сверяются
                       с profile.title_block по порядку: кегль первого run
                       (±0.5), bold, italic (null в профиле = не проверять),
                       выравнивание center. Несоответствие — нарушение
                       с ожидаемым/фактическим. Если Heading-ов нет —
                       нарушение «в документе нет Heading-разделов».
  3. Heading-и       — уровни последовательны: после Heading 1 допустимы 1/2,
                       после Heading 2 — 1/2/3. Скачок 1→3 — предупреждение
                       [WARN]; предупреждения НЕ влияют на exit-код.
  4. Таблицы         — style.name ∈ profile.table_styles[].name; в data-рядах
                       кегль первого run ячейки = cell_size_pt (±0.5); в
                       header-ряде (при header_bold=true) каждый run жирный.
                       Жирность и кегль разрешаются по цепочке: run → стиль
                       абзаца → docDefaults; для жирности в header-ряде
                       дополнительно учитывается условное форматирование
                       firstRow зарегистрированного табличного стиля
                       (эталон набран так, что bold в ячейках не выставлен
                       явно). Нарушения агрегируются по таблице:
                       «table#3: 4 ячейки с кеглем 11.0 вместо 10.0».
  5. Шрифты          — эффективное имя шрифта каждого run (явное на run →
                       цепочка стилей абзаца → docDefaults; нигде не задано —
                       считается наследуемым из темы и НЕ проверяется)
                       должно равняться profile.style_roles.normal.font.
                       Кегль: явный size в Normal-абзаце вне титульного блока
                       должен равняться normal.size_pt ИЛИ совпадать с одним
                       из cell_size_pt зарегистрированных табличных стилей —
                       в эталоне тем же кеглем 10pt набраны примечания под
                       рисунками (несостыковка profile↔эталон решена через
                       уже записанный в профиле cell_size_pt). Таблицы и
                       титул проверяются отдельными проверками.
  6. Страница        — sections[0]: ширина/высота страницы, ориентация
                       и поля равны profile.page (допуск ±1 EMU на
                       округление).

Итог: нарушений (без [WARN]) > 0 → «VISUAL FAIL: N нарушений», exit 1;
иначе «VISUAL PASS (+M предупреждений)», exit 0. Коды 2 — ошибки
использования/IO. Отчёт по-русски в stdout и файл visual_report.txt рядом
с docx (или путь из --out).

Зависимости: python-docx. Python 3.12.
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from collections import Counter
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

ALLOWED_PARA_STYLES = ("Normal", "Heading 1", "Heading 2", "Heading 3",
                       "List Bullet")
HEADING_RE = re.compile(r"^Heading ([123])$")
SIZE_TOL = 0.5      # допуск на кегль, pt
EMU_TOL = 1         # допуск на размеры страницы/поля, EMU

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
ALIGN_NAME = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    None: "не задано (наследуется)",
}


# ------------------------------------------------------------------ helpers

def iter_style_chain(style):
    """Стиль → базовый → … (с защитой от циклов)."""
    seen = set()
    while style is not None and id(style) not in seen:
        seen.add(id(style))
        yield style
        style = style.base_style


def doc_defaults(document):
    """(font_name, size_pt) из w:docDefaults/w:rPrDefault (или None, None)."""
    rpr = document.styles.element.xpath("./w:docDefaults/w:rPrDefault/w:rPr")
    if not rpr:
        return None, None
    fonts = rpr[0].xpath("./w:rFonts/@w:ascii")
    sz = rpr[0].xpath("./w:sz/@w:val")
    return (fonts[0] if fonts else None,
            int(sz[0]) / 2.0 if sz else None)


def effective_font_name(run, para, dd_font):
    """Явное имя на run → цепочка стилей → docDefaults (None = наследуется)."""
    if run.font.name:
        return run.font.name
    for st in iter_style_chain(para.style):
        if st.font.name:
            return st.font.name
    return dd_font


def effective_size_pt(run, para, dd_size):
    """Кегль: явный на run → цепочка стилей → docDefaults (None = неизвестен)."""
    if run.font.size is not None:
        return run.font.size.pt
    for st in iter_style_chain(para.style):
        if st.font.size is not None:
            return st.font.size.pt
    return dd_size


def effective_bold(run, para):
    """Жирность: run → цепочка стилей абзаца (None = определить нельзя)."""
    if run.bold is not None:
        return run.bold
    for st in iter_style_chain(para.style):
        if st.font.bold is not None:
            return st.font.bold
    return None


def effective_italic(run, para):
    if run.italic is not None:
        return run.italic
    for st in iter_style_chain(para.style):
        if st.font.italic is not None:
            return st.font.italic
    return None


def effective_alignment(para):
    """Выравнивание: абзац → цепочка стилей (None = не задано)."""
    if para.alignment is not None:
        return para.alignment
    for st in iter_style_chain(para.style):
        if st.paragraph_format.alignment is not None:
            return st.paragraph_format.alignment
    return None


def table_firstrow_bold(table):
    """Жирность из условного форматирования firstRow табличного стиля."""
    style = table.style
    if style is None:
        return None
    els = style.element.xpath('./w:tblStylePr[@w:type="firstRow"]/w:rPr/w:b')
    if not els:
        return None
    val = els[0].xpath("@w:val")
    if val and str(val[0]).lower() in ("0", "false", "none", "off"):
        return False
    return True


def first_heading_index(paras):
    for i, p in enumerate(paras):
        if HEADING_RE.match(p.style.name or ""):
            return i
    return None


def title_zone(paras):
    """Индексы непустых абзацев до первого Heading (в порядке документа)."""
    cut = first_heading_index(paras)
    if cut is None:
        cut = len(paras)
    return [i for i in range(cut) if paras[i].text.strip()]


def cell_key(cell):
    """Стабильный ключ ячейки (защита от дублей при объединении ячеек)."""
    tc = cell._tc
    return tc.getroottree().getpath(tc)


def fmt_size(v):
    return "не задан" if v is None else f"{round(v, 1)}"


def fmt_bool(v):
    return {True: "True", False: "False", None: "не задано"}[v]


# ------------------------------------------------------------------- checks

def check_para_styles(paras):
    """1. Каждый непустой абзац — стиль из допустимого множества."""
    msgs = []
    for i, p in enumerate(paras):
        if not p.text.strip():
            continue
        name = p.style.name or ""
        if name not in ALLOWED_PARA_STYLES:
            msgs.append(f"[стили] para#{i}: недопустимый стиль '{name}' "
                        f"(допустимо: {', '.join(ALLOWED_PARA_STYLES)})")
    return msgs


def check_title_block(paras, profile, dd_size):
    """2. Титульный блок против profile.title_block."""
    msgs = []
    title_block = profile.get("title_block") or []
    if first_heading_index(paras) is None:
        msgs.append("[титул] в документе нет Heading-разделов")
    zone = title_zone(paras)
    if len(zone) < len(title_block):
        msgs.append(f"[титул] непустых абзацев до первого Heading ({len(zone)}) "
                    f"меньше, чем в title_block ({len(title_block)})")
    if len(zone) > len(title_block):
        msgs.append(f"[титул] непустых абзацев до первого Heading ({len(zone)}) "
                    f"больше, чем в title_block ({len(title_block)})")
    for k, pi in enumerate(zone[:len(title_block)]):
        p = paras[pi]
        exp = title_block[k]
        loc = f"para#{pi} (title_block[{k}])"
        if not p.runs:
            msgs.append(f"[титул] {loc}: абзац без run'ов — сверять нечего")
            continue
        run = p.runs[0]
        exp_size = exp.get("size_pt")
        if exp_size is not None:
            act = effective_size_pt(run, p, dd_size)
            if act is None or abs(act - exp_size) > SIZE_TOL:
                msgs.append(f"[титул] {loc}: кегль первого run "
                            f"{fmt_size(act)} ≠ {exp_size} (±{SIZE_TOL})")
        if exp.get("bold") is not None:
            act = effective_bold(run, p)
            if act is not True:
                msgs.append(f"[титул] {loc}: жирность первого run "
                            f"{fmt_bool(act)} ≠ True")
        if exp.get("italic") is not None:
            act = effective_italic(run, p)
            if act is not True:
                msgs.append(f"[титул] {loc}: курсив первого run "
                            f"{fmt_bool(act)} ≠ True")
        exp_align = exp.get("alignment")
        if exp_align:
            want = ALIGN_MAP.get(exp_align)
            act = effective_alignment(p)
            if act != want:
                msgs.append(f"[титул] {loc}: выравнивание "
                            f"{ALIGN_NAME.get(act, act)} ≠ {exp_align}")
    return msgs


def check_headings(paras):
    """3. Последовательность уровней заголовков (скачок 1→3 — [WARN])."""
    warns = []
    prev = None
    for i, p in enumerate(paras):
        m = HEADING_RE.match(p.style.name or "")
        if not m:
            continue
        lvl = int(m.group(1))
        if prev is not None and prev == 1 and lvl == 3:
            warns.append(f"[WARN] [заголовки] скачок уровня Heading 1 → "
                         f"Heading 3 — para#{i}")
        prev = lvl
    return warns


def check_tables(document, profile, dd_size):
    """4. Стили таблиц, кегль data-ячеек, жирность header-ряда."""
    msgs = []
    registered = {ts.get("name"): ts
                  for ts in profile.get("table_styles") or [] if ts.get("name")}
    default_entry = next(iter(registered.values()), None)
    for ti, table in enumerate(document.tables):
        loc = f"table#{ti}"
        entry = None
        try:
            style_name = table.style.name if table.style is not None else None
        except KeyError:
            style_name = None
        if style_name is None:
            msgs.append(f"[таблицы] {loc}: таблица без стиля "
                        f"(w:tblStyle отсутствует)")
        elif style_name not in registered:
            msgs.append(f"[таблицы] {loc}: стиль таблицы '{style_name}' не "
                        f"зарегистрирован в profile.table_styles "
                        f"(допустимо: {', '.join(registered) or '—'})")
        else:
            entry = registered[style_name]
        cell_size = (entry or default_entry or {}).get("cell_size_pt")

        # --- кегль data-ячеек (агрегированно по таблице) -------------------
        if cell_size is not None:
            bad = Counter()
            seen = set()
            for row in table.rows[1:]:
                for cell in row.cells:
                    key = cell_key(cell)
                    if key in seen:
                        continue
                    seen.add(key)
                    para0 = cell.paragraphs[0]
                    if not para0.runs:
                        continue
                    act = effective_size_pt(para0.runs[0], para0, dd_size)
                    if act is None:
                        continue
                    if abs(act - cell_size) > SIZE_TOL:
                        bad[round(act, 1)] += 1
            for act_sz, n in sorted(bad.items()):
                msgs.append(f"[таблицы] {loc}: {n} ячеек с кеглем {act_sz} "
                            f"вместо {cell_size}")

        # --- жирность header-ряда ------------------------------------------
        if entry is None or entry.get("header_bold"):
            style_bold = table_firstrow_bold(table)
            n_bad = 0
            seen = set()
            for cell in table.rows[0].cells:
                key = cell_key(cell)
                if key in seen:
                    continue
                seen.add(key)
                for p in cell.paragraphs:
                    for r in p.runs:
                        b = effective_bold(r, p)
                        if b is None:
                            b = style_bold
                        if b is not True:
                            n_bad += 1
            if n_bad:
                msgs.append(f"[таблицы] {loc}: {n_bad} run(ов) в header-ряде "
                            f"без жирного начертания")
    return msgs


def check_fonts(paras, profile, dd_font, dd_size, title_paras):
    """5. Эффективные имена шрифтов и явные кегли в Normal-абзацах."""
    msgs = []
    roles = profile.get("style_roles") or {}
    normal_font = (roles.get("normal") or {}).get("font")
    normal_size = (roles.get("normal") or {}).get("size_pt")
    allowed_sizes = ([normal_size] if normal_size is not None else []) + [
        ts.get("cell_size_pt") for ts in profile.get("table_styles") or []
        if ts.get("cell_size_pt") is not None]
    title_ids = {id(p) for p in title_paras}
    for i, p in enumerate(paras):
        if not p.text.strip():
            continue
        for k, r in enumerate(p.runs):
            name = effective_font_name(r, p, dd_font)
            if name and normal_font and name != normal_font:
                msgs.append(f"[шрифты] para#{i} run#{k}: шрифт '{name}' "
                            f"вместо '{normal_font}'")
        if p.style.name == "Normal" and id(p) not in title_ids:
            for k, r in enumerate(p.runs):
                if r.font.size is None:
                    continue  # проверяются только явные отклонения
                act = r.font.size.pt
                if normal_size is None:
                    continue
                ok = abs(act - normal_size) <= SIZE_TOL or any(
                    abs(act - s) <= SIZE_TOL for s in allowed_sizes)
                if not ok:
                    allowed = ", ".join(str(round(s, 1)) for s in allowed_sizes)
                    msgs.append(f"[кегль] para#{i} run#{k}: явный кегль "
                                f"{round(act, 1)} в Normal-абзаце "
                                f"(допустимо: {allowed})")
    return msgs


def check_page(document, profile):
    """6. Параметры страницы sections[0] против profile.page (±1 EMU)."""
    msgs = []
    page = profile.get("page") or {}
    margins = page.get("margins_emu") or {}
    sec = document.sections[0]
    pairs = [
        ("ширина страницы", sec.page_width, page.get("width_emu")),
        ("высота страницы", sec.page_height, page.get("height_emu")),
        ("поле top", sec.top_margin, margins.get("top")),
        ("поле right", sec.right_margin, margins.get("right")),
        ("поле bottom", sec.bottom_margin, margins.get("bottom")),
        ("поле left", sec.left_margin, margins.get("left")),
    ]
    for label, actual, expected in pairs:
        if expected is None:
            continue
        if actual is None or abs(int(actual) - int(expected)) > EMU_TOL:
            act_s = "не задано" if actual is None else str(int(actual))
            msgs.append(f"[страница] {label} {act_s} EMU ≠ {expected} EMU "
                        f"(±{EMU_TOL})")
    want_orientation = page.get("orientation")
    if want_orientation:
        actual_landscape = sec.orientation is not None and \
            int(sec.orientation) == 1
        expected_landscape = want_orientation == "landscape"
        if actual_landscape != expected_landscape:
            msgs.append(f"[страница] ориентация "
                        f"{'landscape' if actual_landscape else 'portrait'} "
                        f"≠ {want_orientation}")
    return msgs


# ------------------------------------------------------------------- report

def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(
        description="Визуальный линт docx против стилевого профиля.")
    ap.add_argument("--docx", required=True, help="собранный .docx")
    ap.add_argument("--profile", required=True, help="style-profile.json")
    ap.add_argument("--out", help="путь отчёта (по умолчанию "
                                  "visual_report.txt рядом с docx)")
    args = ap.parse_args(argv)

    target = Path(args.docx)
    if not target.is_file():
        print(f"[ошибка] docx не найден: {target}", file=sys.stderr)
        return 2
    profile_path = Path(args.profile)
    if not profile_path.is_file():
        print(f"[ошибка] профиль не найден: {profile_path}", file=sys.stderr)
        return 2
    try:
        profile = json.loads(profile_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        print(f"[ошибка] профиль — некорректный JSON: {exc}", file=sys.stderr)
        return 2

    try:
        d = docx.Document(str(target))
    except Exception as exc:  # noqa: BLE001 — повреждённый docx
        print(f"[ошибка] docx не открывается: {exc}", file=sys.stderr)
        return 2

    paras = d.paragraphs
    dd_font, dd_size = doc_defaults(d)
    title_paras = [paras[i] for i in title_zone(paras)][
        :len(profile.get("title_block") or [])]

    violations: list[str] = []
    violations += check_para_styles(paras)
    violations += check_title_block(paras, profile, dd_size)
    warns: list[str] = list(check_headings(paras))
    violations += check_tables(d, profile, dd_size)
    violations += check_fonts(paras, profile, dd_font, dd_size, title_paras)
    violations += check_page(d, profile)

    lines = [
        "=" * 72,
        f"lint_visual: {target}",
        f"Профиль: {profile_path} | Абзацев: {len(paras)} | "
        f"Таблиц: {len(d.tables)} | Разделов (sectPr): {len(d.sections)}",
        "=" * 72,
        f"НАРУШЕНИЯ (FAIL): {len(violations)}",
    ]
    for n, m in enumerate(violations, start=1):
        lines.append(f"  {n}. {m}")
    lines.append(f"ПРЕДУПРЕЖДЕНИЯ (warn): {len(warns)}")
    for n, w in enumerate(warns, start=1):
        lines.append(f"  {n}. {w}")

    report_path = Path(args.out) if args.out \
        else target.resolve().parent / "visual_report.txt"
    report_path.parent.mkdir(parents=True, exist_ok=True)
    if violations:
        lines.append(f"ИТОГ: VISUAL FAIL: {len(violations)} нарушений")
    else:
        lines.append(f"ИТОГ: VISUAL PASS (+{len(warns)} предупреждений)")
    report_path.write_text("\n".join(lines) + "\n", encoding="utf-8")

    print("\n".join(lines))
    return 1 if violations else 0


if __name__ == "__main__":
    sys.exit(main())
