#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
extract_profile.py — Task 4 (Steps 4.1–4.6) комплекта report-style-kit.

Детерминированно извлекает профиль стиля из образцового .docx-отчёта:
  Step 4.1  page, style_roles (с наследованием/docDefaults/темой), table_styles,
            microtypography (счётчики паттернов + детерминанты + примеры),
            glossary_terms (латинские токены, POT_N, имена файлов)
  Step 4.2  title_block, block_patterns (подписи рисунков, формулы, lead-in)
  Step 4.3  table_templates (кластеризация по сигнатуре header-строки)
            + section_blueprint (обход body в порядке документа)
  Step 4.4  template.docx — копия образца без содержимого body (стили/тема/sectPr нетронуты)
  Step 4.5  work/styleguide.draft.md + work/glossary.draft.md
  Step 4.6  work/slang-blacklist.txt — копия базы; категории, с которыми
            пересекается глоссарий образца, комментируются «# легально в образце:»

CLI:
    python3 extract_profile.py --sample X.docx --outdir work/

Зависимости: python-docx (остальное — стандартная библиотека). Python 3.12.
"""
from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
import xml.etree.ElementTree as ET
import zipfile
from collections import Counter, OrderedDict
from pathlib import Path

import docx
from docx.oxml.ns import qn

A_NS = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
TWIPS_TO_EMU = 635
HEADING_STYLES = {"Heading 1": 1, "Heading 2": 2, "Heading 3": 3}
BLOCK_TYPES = ("thesis", "lead_in", "formula", "bullets", "table", "figure", "verdict")

TOKEN_RE = re.compile(r"[A-Za-z_][A-Za-z0-9_\-\.]{1,}")
POT_RE = re.compile(r"\bPOT_\d+\b")
NUM_PREFIX_RE = re.compile(r"^\d+(?:\.\d+)*\.?\s*")
# Абсолютные пути хоста, встретившиеся в самих текстах образца: в артефакты
# комплекта (примеры-цитаты) попадают только в замаскированном виде —
# требование автономности закрытого контура (аудит copy-out).
ABS_PATH_RE = re.compile(r"(?<![\w])/(?:root|home|Users)/[^\s|»«()]*")


def sanitize_sample_text(text: str) -> str:
    """Абсолютные пути из текста образца → «<путь-образца>» (только цитаты-примеры)."""
    return ABS_PATH_RE.sub("<путь-образца>", text)
FIGURE_RE_DEFAULT = r"^Рис\. \d+\."
FIGURE_RE_FALLBACK = r"^Рисунок \d+"

# Паттерны микротипографики (Step 4.1): имя -> regex. Счёт по каждому тексту отдельно
# (anchors ^ должны работать на границах абзацев/ячеек, а не всего документа).
MINUS_ASCII_RE = re.compile(r"\d-\d|(?:^|[\s(])-\d")
RANGE_ASCII_RE = re.compile(r"\d-\d")
PATTERNS = OrderedDict([
    ("decimal_point", re.compile(r"\d\.\d")),
    ("decimal_comma", re.compile(r"\d,\d")),
    ("minus_unicode", re.compile("\u2212")),
    ("minus_ascii_hyphen", MINUS_ASCII_RE),
    ("percent_spaced", re.compile(r"\d\s%")),
    ("percent_glued", re.compile(r"\d%")),
    ("range_endash", re.compile(r"\d–\d")),
    ("quotes_guillemets", re.compile("«")),
    ("quotes_straight", re.compile("\u0022")),
    ("arrow", re.compile("→")),
    ("approx", re.compile("≈")),
    ("multiplication", re.compile("×")),
])


# --------------------------------------------------------------------------- helpers

def load_theme_fonts(docx_path: Path) -> dict:
    """{'major': <typeface>, 'minor': <typeface>} из word/theme/theme1.xml."""
    fonts = {}
    try:
        with zipfile.ZipFile(docx_path) as z:
            parts = [n for n in z.namelist() if n.startswith("word/theme/")]
            if not parts:
                return fonts
            root = ET.fromstring(z.read(parts[0]))
            for tag, key in (("majorFont", "major"), ("minorFont", "minor")):
                el = root.find(f".//{A_NS}fontScheme/{A_NS}{tag}/{A_NS}latin")
                if el is not None:
                    fonts[key] = el.get("typeface")
    except (KeyError, ET.ParseError) as exc:
        print(f"[warn] тема не прочитана: {exc}", file=sys.stderr)
    return fonts


def theme_font_for(theme_attr: str | None, theme_fonts: dict) -> str | None:
    if not theme_attr:
        return None
    t = theme_attr.lower()
    if t.startswith("major"):
        return theme_fonts.get("major")
    if t.startswith("minor"):
        return theme_fonts.get("minor")
    return None


def doc_defaults_rfonts(styles_el):
    """XPath w:docDefaults/w:rPrDefault/w:rPr/w:rFonts."""
    dd = styles_el.find(qn("w:docDefaults"))
    if dd is None:
        return None
    rpr = dd.find(f"{qn('w:rPrDefault')}/{qn('w:rPr')}")
    if rpr is None:
        return None
    return rpr.find(qn("w:rFonts"))


def _own_bool(rpr, tag: str) -> bool | None:
    """Значение w:b/w:i собственного rPr стиля: True/False/None (нет элемента)."""
    if rpr is None:
        return None
    el = rpr.find(qn(tag))
    if el is None:
        return None
    val = el.get(qn("w:val"))
    if val is not None and val.strip().lower() in ("0", "false", "off"):
        return False
    return True


def _own_spacing(style_el) -> tuple[int | None, int | None]:
    """spacing before/after (twips -> EMU) из собственного pPr стиля."""
    ppr = style_el.find(qn("w:pPr"))
    if ppr is None:
        return None, None
    sp = ppr.find(qn("w:spacing"))
    if sp is None:
        return None, None
    before = after = None
    if sp.get(qn("w:before")) is not None:
        before = int(sp.get(qn("w:before"))) * TWIPS_TO_EMU
    if sp.get(qn("w:after")) is not None:
        after = int(sp.get(qn("w:after"))) * TWIPS_TO_EMU
    return before, after


def resolve_font(style, styles_el, theme_fonts: dict) -> tuple[str | None, str | None]:
    """Эффективный шрифт стиля: own ascii -> own asciiTheme -> базовая цепочка -> docDefaults."""
    rpr = style.element.find(qn("w:rPr"))
    if rpr is not None:
        rf = rpr.find(qn("w:rFonts"))
        if rf is not None:
            if rf.get(qn("w:ascii")):
                return rf.get(qn("w:ascii")), "style"
            if rf.get(qn("w:asciiTheme")):
                return theme_font_for(rf.get(qn("w:asciiTheme")), theme_fonts), "theme"
    base, seen = style.base_style, {style.style_id}
    while base is not None and base.style_id not in seen:
        seen.add(base.style_id)
        brpr = base.element.find(qn("w:rPr"))
        brf = brpr.find(qn("w:rFonts")) if brpr is not None else None
        if brf is not None:
            if brf.get(qn("w:ascii")):
                return brf.get(qn("w:ascii")), "inherited"
            if brf.get(qn("w:asciiTheme")):
                return theme_font_for(brf.get(qn("w:asciiTheme")), theme_fonts), "inherited"
        base = base.base_style
    rf = doc_defaults_rfonts(styles_el)
    if rf is not None:
        if rf.get(qn("w:ascii")):
            return rf.get(qn("w:ascii")), "docDefaults"
        if rf.get(qn("w:asciiTheme")):
            return theme_font_for(rf.get(qn("w:asciiTheme")), theme_fonts), "docDefaults"
    return None, None


def style_role(style, styles_el, theme_fonts: dict) -> dict:
    """Роль стиля: font — с наследованием; size/bold/color — собственные значения стиля."""
    rpr = style.element.find(qn("w:rPr"))
    font, source = resolve_font(style, styles_el, theme_fonts)
    size_pt = None
    if rpr is not None:
        sz = rpr.find(qn("w:sz"))
        if sz is not None and sz.get(qn("w:val")):
            size_pt = int(sz.get(qn("w:val"))) / 2.0
    color_hex = None
    if rpr is not None:
        col = rpr.find(qn("w:color"))
        if col is not None and col.get(qn("w:val")):
            color_hex = col.get(qn("w:val"))
    before, after = _own_spacing(style.element)
    return {
        "font": font,
        "size_pt": size_pt,
        "bold": _own_bool(rpr, "w:b"),
        "color_hex": color_hex,
        "spacing_before_emu": before,
        "spacing_after_emu": after,
        "source": source,
    }


def first_nonempty_run(cell):
    """Первый непустой run в ячейке (по всем абзацам)."""
    for p in cell.paragraphs:
        for r in p.runs:
            if r.text.strip():
                return r
    return None


def para_has_image(paragraph) -> bool:
    return bool(paragraph._p.findall(".//" + qn("a:blip")))


def is_fully_italic(paragraph) -> bool:
    """Непустой абзац, у которого все текстовые run'ы курсивные."""
    if not paragraph.text.strip():
        return False
    runs = [r for r in paragraph.runs if r.text.strip()]
    if not runs:
        return False
    return all(r.italic is True for r in runs)


def is_lead_in(paragraph, terminator: str = ":") -> bool:
    """Абзац, начинающийся с жирного run'а, оканчивающегося на терминатор."""
    runs = paragraph.runs
    if not runs:
        return False
    r0 = runs[0]
    return bool(r0.bold) and r0.text.rstrip().endswith(terminator)


def strip_numbering(text: str) -> str:
    return NUM_PREFIX_RE.sub("", text).strip()


def first_sentence_with(text: str, term: str, limit: int) -> str | None:
    """Первое предложение текста, содержащее термин (обрезанное до limit)."""
    low = text.lower()
    if term not in low:
        return None
    for sent in re.split(r"(?<=[.!?…])\s+", text.strip()):
        if term in sent.lower():
            return sent[:limit]
    return text.strip()[:limit]


def ordered_texts(doc) -> list[str]:
    """Все непустые тексты в порядке документа: абзацы body + ячейки таблиц."""
    texts = []
    para_map = {p._p: p for p in doc.paragraphs}
    for child in doc.element.body:
        if child.tag == qn("w:p"):
            p = para_map.get(child)
            if p is not None and p.text.strip():
                texts.append(p.text)
        elif child.tag == qn("w:tbl"):
            for t in doc.tables:
                if t._tbl is child:
                    for row in t.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                texts.append(cell.text)
                    break
    return texts


# --------------------------------------------------------------------------- steps

def extract_page(doc) -> dict:
    s = doc.sections[0]
    return {
        "width_emu": int(s.page_width),
        "height_emu": int(s.page_height),
        "orientation": s.orientation.name.lower(),
        "margins_emu": {
            "top": int(s.top_margin),
            "right": int(s.right_margin),
            "bottom": int(s.bottom_margin),
            "left": int(s.left_margin),
        },
    }


def extract_style_roles(doc, theme_fonts: dict) -> dict:
    styles_el = doc.styles.element
    roles = {}
    for key, sname in (("normal", "Normal"), ("heading1", "Heading 1"),
                       ("heading2", "Heading 2"), ("heading3", "Heading 3"),
                       ("list_bullet", "List Bullet")):
        try:
            roles[key] = style_role(doc.styles[sname], styles_el, theme_fonts)
        except KeyError:
            print(f"[stop] стиль {sname!r} не найден в образце", file=sys.stderr)
            raise SystemExit(3)
    return roles


def extract_table_styles(doc, normal_size: float | None) -> list[dict]:
    out, seen = [], set()
    for t in doc.tables:
        sname = t.style.name if t.style is not None else "Normal"
        if sname in seen:
            continue
        seen.add(sname)
        rows = t.rows
        data_row = rows[1] if len(rows) > 1 else (rows[0] if rows else None)
        run = None
        if data_row is not None:
            for c in data_row.cells:
                run = first_nonempty_run(c)
                if run is not None:
                    break
        hrun = None
        if rows:
            for c in rows[0].cells:
                hrun = first_nonempty_run(c)
                if hrun is not None:
                    break
        cell_size = None
        if run is not None and run.font.size is not None:
            cell_size = run.font.size.pt
        if cell_size is None and hrun is not None and hrun.font.size is not None:
            cell_size = hrun.font.size.pt
        if cell_size is None:
            cell_size = normal_size if normal_size is not None else 11.0
        out.append({
            "name": sname,
            "cell_size_pt": float(cell_size),
            "header_bold": bool(hrun is not None and hrun.bold),
        })
    return out


def scan_microtypography(texts: list[str]) -> dict:
    counts, examples = OrderedDict(), {}
    for name, rx in PATTERNS.items():
        total = 0
        for t in texts:
            total += len(rx.findall(t))
        counts[name] = total
    for name, rx in PATTERNS.items():
        for t in texts:
            m = rx.search(t)
        if m:
            start = max(0, m.start() - 8)
            examples[name] = sanitize_sample_text(t[start:start + 40])
            break

    def pick(a: str, b: str, va: str, vb: str, default: str) -> str:
        ca, cb = counts[a], counts[b]
        if ca and cb:
            return "mixed"
        if ca:
            return va
        if cb:
            return vb
        return default

    micro = {
        "decimal_separator": pick("decimal_point", "decimal_comma", "point", "comma", "point"),
        "minus_char": pick("minus_unicode", "minus_ascii_hyphen", "U+2212", "ascii_hyphen", "U+2212"),
        "percent_spacing": pick("percent_spaced", "percent_glued", "spaced", "glued", "spaced"),
        "range_dash": "en_dash" if counts["range_endash"] else (
            "hyphen" if RANGE_ASCII_RE and any(RANGE_ASCII_RE.findall(t) for t in texts) else "en_dash"),
        "quotes": pick("quotes_guillemets", "quotes_straight", "guillemets", "straight", "guillemets"),
        "arrow": counts["arrow"] > 0,
        "approx": counts["approx"] > 0,
        "multiplication": counts["multiplication"] > 0,
        "patterns": dict(counts),
        "examples": examples,
    }
    # range_dash "mixed": есть и en-dash, и ascii-диапазоны \d-\d
    if counts["range_endash"] and any(RANGE_ASCII_RE.search(t) for t in texts):
        micro["range_dash"] = "mixed"
    return micro


def build_glossary(texts: list[str], top: int = 60) -> list[dict]:
    counter: Counter = Counter()
    for t in texts:
        for m in TOKEN_RE.finditer(t):
            tok = m.group(0).rstrip(".").lower()
            if len(tok) >= 2:
                counter[tok] += 1
    pot_ids, pot_total = set(), 0
    for t in texts:
        for m in POT_RE.finditer(t):
            pot_ids.add(m.group(0))
            pot_total += 1
    if len(pot_ids) >= 3:
        counter["pot_n"] += pot_total
    result = []
    for tok, cnt in counter.most_common(top):
        result.append({"term": "POT_N" if tok == "pot_n" else tok, "count": cnt})
    return result


def extract_title_block(doc, fallback_size: float | None) -> list[dict]:
    block = []
    for p in doc.paragraphs:
        if p.style.name in HEADING_STYLES:
            break
        if not p.text.strip():
            continue
        r0 = p.runs[0] if p.runs else None
        size = r0.font.size.pt if r0 is not None and r0.font.size is not None else None
        if size is None:
            size = fallback_size if fallback_size is not None else 11.0
        block.append({
            "size_pt": float(size),
            "bold": r0.bold if r0 is not None else None,
            "italic": r0.italic if r0 is not None else None,
            "alignment": p.alignment.name.lower() if p.alignment is not None else None,
        })
    return block


def extract_block_patterns(doc, texts: list[str]) -> dict:
    caption_re = FIGURE_RE_DEFAULT
    caption_count = sum(1 for t in texts if re.search(caption_re, t))
    if caption_count == 0:
        caption_re = FIGURE_RE_FALLBACK
        caption_count = sum(1 for t in texts if re.search(caption_re, t))
    formula_count = 0
    lead_in_count = 0
    for p in doc.paragraphs:
        if p.style.name == "Normal" and is_fully_italic(p):
            formula_count += 1
        if is_lead_in(p):
            lead_in_count += 1
    return {
        "figure_caption_regex": caption_re,
        "formula_style": "italic" if formula_count >= 3 else "none",
        "lead_in_terminator": ":",
        "lead_in_count": lead_in_count,
        "figure_caption_count": caption_count,
        "formula_count": formula_count,
    }


def template_name(header: list[str]) -> str | None:
    """Эвристики именования (порядок подобран так, чтобы эталон получал
    per_pot_params/version_history вместо более ранних общих правил)."""
    h = [c.strip().lower() for c in header]
    joined = "|".join(h)
    first = h[0] if h else ""
    if first == "pot" or first.startswith("pot"):
        return "per_pot_params"
    if "версия" in joined:
        return "version_history"
    if "гипотеза" in joined:
        return "hypotheses_verdicts"
    if "риск" in joined:
        return "risks_register"
    if "mae" in joined or any(d in joined for d in ("δ", "∆", "∂")):
        return "metrics_comparison"
    if len(h) == 2 and any(k in first for k in ("параметр", "период", "версия", "framework")):
        return "key_value_params"
    return None


def extract_table_templates(doc) -> list[dict]:
    groups: "OrderedDict[str, dict]" = OrderedDict()
    for idx, t in enumerate(doc.tables):
        header = [c.text.strip() for c in t.rows[0].cells] if t.rows else []
        sig = "|".join(h.strip().lower() for h in header)
        g = groups.setdefault(sig, {"header": header, "indices": []})
        g["indices"].append(idx)
    templates, used_names, auto_n = [], Counter(), 0
    for g in groups.values():
        name = template_name(g["header"])
        if name is None:
            auto_n += 1
            name = f"table_template_{auto_n}"
        used_names[name] += 1
        if used_names[name] > 1:
            name = f"{name}_{used_names[name]}"
        templates.append({
            "name": name,
            "header": g["header"],
            "exemplar_table_index": g["indices"][0],
            "occurrences": len(g["indices"]),
        })
    return templates


def extract_section_blueprint(doc) -> list[dict]:
    para_map = {p._p: p for p in doc.paragraphs}
    table_map = {t._tbl: t for t in doc.tables}
    blueprint: list[dict] = []
    current: dict | None = None

    def push(block: str):
        # Блоки до первого заголовка (титульный блок) в blueprint не входят.
        if current is None:
            return
        if block == "bullets" and current["block_sequence"][-1:] == ["bullets"]:
            return  # подряд идущие List Bullet -> один блок bullets
        current["block_sequence"].append(block)

    for child in doc.element.body:
        if child.tag == qn("w:p"):
            p = para_map.get(child)
            if p is None:
                continue
            sname = p.style.name
            if sname in HEADING_STYLES:
                current = {
                    "level": HEADING_STYLES[sname],
                    "heading": strip_numbering(p.text),
                    "block_sequence": [],
                }
                blueprint.append(current)
                continue
            if para_has_image(p):
                push("figure")
                continue
            if not p.text.strip():
                continue  # пустые абзацы-разделители не блоки
            if sname == "List Bullet":
                push("bullets")
            elif is_lead_in(p):
                push("lead_in")
            elif p.style.name == "Normal" and is_fully_italic(p):
                push("formula")
            else:
                push("thesis")
        elif child.tag == qn("w:tbl"):
            if child in table_map:
                push("table")
    return blueprint


def build_template_docx(sample: Path, outdir: Path) -> Path:
    """Step 4.4: копия образца, body очищен (кроме sectPr)."""
    dst = outdir / "template.docx"
    shutil.copy2(sample, dst)
    d = docx.Document(str(dst))
    body = d.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)
    d.save(str(dst))
    return dst


# --------------------------------------------------------------------------- drafts

def _md_table(rows: list[list[str]]) -> str:
    out = ["| " + " | ".join(rows[0]) + " |", "|" + "|".join(["---"] * len(rows[0])) + "|"]
    for r in rows[1:]:
        out.append("| " + " | ".join(c.replace("|", "\\|") for c in r) + " |")
    return "\n".join(out)


def write_styleguide_draft(path: Path, profile: dict, sample: Path, texts: list[str],
                           examples: dict, block_examples: dict, first_person: int) -> None:
    micro = profile["microtypography"]
    det_rows = [["Параметр", "Значение", "Пример из образца"]]
    for key, label in (("decimal_separator", "Десятичный разделитель"),
                       ("minus_char", "Минус"),
                       ("percent_spacing", "Процент"),
                       ("range_dash", "Диапазон"),
                       ("quotes", "Кавычки")):
        ex_key = {"decimal_separator": "decimal_point", "minus_char": "minus_unicode",
                  "percent_spacing": "percent_spaced", "range_dash": "range_endash",
                  "quotes": "quotes_guillemets"}[key]
        ex = sanitize_sample_text(examples.get(ex_key, "—"))
        det_rows.append([label, str(micro[key]), ex])
    for key, label, pname in (("arrow", "Стрелка →", "arrow"),
                              ("approx", "Примерно ≈", "approx"),
                              ("multiplication", "Умножение ×", "multiplication")):
        det_rows.append([label, "да" if micro[key] else "нет",
                         sanitize_sample_text(examples.get(pname, "—"))])

    tone = ("Первое лицо не используется (вхождений 1-го лица: {n}). "
            "Отчёт — безличная констатация фактов.").format(n=first_person) \
        if first_person == 0 else \
        "Обнаружены вхождения 1-го лица ({n}) — проверить образец вручную.".format(n=first_person)

    tpl_rows = [["Шаблон", "Заголовки"]]
    for t in profile["table_templates"]:
        tpl_rows.append([t["name"], " | ".join(t["header"])])
    if len(tpl_rows) == 1:
        tpl_rows.append(["—", "—"])

    legend = [
        ("thesis", "Тезис: абзац 2–5 предложений, вывод — первым предложением."),
        ("lead_in", "Подводка: жирный заголовок-фраза с «:», далее пояснение."),
        ("formula", "Формульная строка: отдельный полностью курсивный абзац."),
        ("bullets", "Маркированный список (стиль List Bullet)."),
        ("table", "Таблица зарегистрированным стилем, header bold, ячейки 10pt."),
        ("figure", "Рисунок: подпись «Рис. N. …» + изображение + пояснение."),
        ("verdict", "Вердикт по гипотезе: ПОДТВЕРЖДЕНА / ОТВЕРГНУТА и т.п."),
    ]
    block_rows = [["Тип", "Легенда", "Пример из образца (≤200 симв.)"]]
    for btype, desc in legend:
        ex = block_examples.get(btype)
        block_rows.append([btype, desc,
                           sanitize_sample_text((ex or "— нет в образце —"))[:200]])

    lines = [
        f"# Styleguide (draft) — {sample.name}",
        "",
        "> Черновик сгенерирован автоматически — курайте перед использованием.",
        "",
        "## Тон",
        "",
        tone,
        "",
        "## Структура абзацев",
        "",
        "Тезис-первое-предложение: абзац открывается выводом, затем контекст и числа.",
        "Пример тезис-абзаца из образца:",
        "",
        "> " + sanitize_sample_text(block_examples.get("thesis", "—") or "—")[:200],
        "",
        "## Числа и типографика",
        "",
        _md_table(det_rows),
        "",
        "Ссылочная строка типографики: минус U+2212, разделитель — точка, "
        f"проценты — {micro['percent_spacing']}, кавычки — {micro['quotes']}.",
        "",
        "Характерный фрагмент образца (минус U+2212 с числом и процентом):",
        "",
    ]
    minus_pct = None
    rx = re.compile(".{0,12}\u2212\\d+(?:[.,]\\d+)?\\s?%.{0,24}")
    for t in texts:
        m = rx.search(t)
        if m:
            minus_pct = m.group(0)
            break
    lines.append("«" + (sanitize_sample_text(minus_pct[:40]) if minus_pct else "— не найдено —") + "»")
    lines += [
        "",
        "## Блоки",
        "",
        _md_table(block_rows),
        "",
        "## Таблицы",
        "",
        _md_table(tpl_rows),
        "",
        "",
    ]
    path.write_text("\n".join(lines), encoding="utf-8")


def write_glossary_draft(path: Path, sample: Path, glossary: list[dict],
                         texts: list[str], top: int = 40) -> None:
    rows = [["Термин", "Частота", "Пример употребления"]]
    for g in glossary[:top]:
        term = g["term"]
        example = None
        for t in texts:
            example = first_sentence_with(t, term, 120)
            if example:
                break
        rows.append([term, str(g["count"]),
                     sanitize_sample_text(example) if example else "—"])
    lines = [
        f"# Glossary (draft) — {sample.name}",
        "",
        "> Черновик сгенерирован автоматически — курайте перед использованием.",
        "",
        _md_table(rows),
        "",
        "",
    ]
    path.write_text("\n".join(lines), encoding="utf-8")


def copy_blacklist(base: Path, dst: Path, glossary: list[dict]) -> list[str]:
    """Step 4.6: копия базы; категории, совпавшие с глоссарием образца,
    комментируются с пометкой «# легально в образце:»."""
    out, changed = [], []
    terms = [g["term"] for g in glossary]
    for line in base.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if stripped and not stripped.startswith("#") and "\t" in line:
            cat, rx = line.split("\t", 1)
            try:
                cre = re.compile(rx.strip(), re.IGNORECASE)
            except re.error:
                out.append(line)
                continue
            matched = [t for t in terms if cre.search(t)]
            if matched:
                out.append(f"# легально в образце: {', '.join(matched)}")
                out.append("# " + line)
                changed.append(cat.strip())
                continue
        out.append(line)
    dst.write_text("\n".join(out) + "\n", encoding="utf-8")
    return changed


# --------------------------------------------------------------------------- validation

def validate_profile(p: dict) -> list[str]:
    """Мини-валидатор против style-profile.schema.json (required + enum'ы)."""
    errs = []
    for key in ("page", "style_roles", "table_styles", "microtypography",
                "table_templates", "section_blueprint"):
        if key not in p:
            errs.append(f"отсутствует обязательное поле: {key}")
    if "page" in p:
        for key in ("width_emu", "height_emu", "orientation", "margins_emu"):
            if key not in p["page"]:
                errs.append(f"page.{key} отсутствует")
        for key, val in (p.get("page", {}).get("margins_emu") or {}).items():
            if not isinstance(val, int):
                errs.append(f"page.margins_emu.{key} не integer")
        if p.get("page", {}).get("orientation") not in ("portrait", "landscape"):
            errs.append("page.orientation вне enum")
    roles = ("normal", "heading1", "heading2", "heading3", "list_bullet")
    for role in roles:
        r = p.get("style_roles", {}).get(role)
        if not isinstance(r, dict):
            errs.append(f"style_roles.{role} отсутствует")
            continue
        for key in ("font", "size_pt", "bold", "color_hex",
                    "spacing_before_emu", "spacing_after_emu"):
            if key not in r:
                errs.append(f"style_roles.{role}.{key} отсутствует")
        ch = r.get("color_hex")
        if ch is not None and not re.fullmatch(r"[0-9A-Fa-f]{6}", ch):
            errs.append(f"style_roles.{role}.color_hex не 6 hex-символов")
    for i, ts in enumerate(p.get("table_styles", [])):
        for key in ("name", "cell_size_pt", "header_bold"):
            if key not in ts:
                errs.append(f"table_styles[{i}].{key} отсутствует")
    for i, tb in enumerate(p.get("title_block", [])):
        for key in ("size_pt", "bold", "italic", "alignment"):
            if key not in tb:
                errs.append(f"title_block[{i}].{key} отсутствует")
    bp = p.get("block_patterns", {})
    for key in ("figure_caption_regex", "formula_style", "lead_in_terminator"):
        if key not in bp:
            errs.append(f"block_patterns.{key} отсутствует")
    micro = p.get("microtypography", {})
    enums = {
        "decimal_separator": {"point", "comma", "mixed"},
        "minus_char": {"U+2212", "ascii_hyphen", "mixed"},
        "percent_spacing": {"spaced", "tight", "glued", "mixed"},
        "range_dash": {"en_dash", "hyphen", "mixed"},
        "quotes": {"guillemets", "straight", "mixed"},
    }
    for key, allowed in enums.items():
        if micro.get(key) not in allowed:
            errs.append(f"microtypography.{key}={micro.get(key)!r} вне enum")
    for key in ("arrow", "approx", "multiplication"):
        if not isinstance(micro.get(key), bool):
            errs.append(f"microtypography.{key} не boolean")
    for i, tt in enumerate(p.get("table_templates", [])):
        for key in ("name", "header", "exemplar_table_index"):
            if key not in tt:
                errs.append(f"table_templates[{i}].{key} отсутствует")
    for i, sb in enumerate(p.get("section_blueprint", [])):
        if sb.get("level") not in (1, 2, 3):
            errs.append(f"section_blueprint[{i}].level вне [1,2,3]")
        for key in ("heading", "block_sequence"):
            if key not in sb:
                errs.append(f"section_blueprint[{i}].{key} отсутствует")
        for b in sb.get("block_sequence", []):
            if b not in BLOCK_TYPES:
                errs.append(f"section_blueprint[{i}] неизвестный тип блока {b!r}")
    for i, g in enumerate(p.get("glossary_terms", [])):
        for key in ("term", "count"):
            if key not in g:
                errs.append(f"glossary_terms[{i}].{key} отсутствует")
    return errs


# --------------------------------------------------------------------------- main

def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description="Извлечение профиля стиля из образцового docx.")
    ap.add_argument("--sample", required=True, help="образцовый .docx отчёт")
    ap.add_argument("--outdir", required=True, help="каталог для артефактов (work/)")
    args = ap.parse_args(argv)

    sample = Path(args.sample).resolve()
    outdir = Path(args.outdir)
    outdir.mkdir(parents=True, exist_ok=True)
    if not sample.is_file():
        print(f"[stop] образец не найден: {sample}", file=sys.stderr)
        return 2

    kit_root = Path(__file__).resolve().parent.parent
    base_blacklist = kit_root / "base" / "slang-blacklist.txt"

    doc = docx.Document(str(sample))
    theme_fonts = load_theme_fonts(sample)
    texts = ordered_texts(doc)
    roles = extract_style_roles(doc, theme_fonts)
    normal_size = roles["normal"]["size_pt"]

    # Steps 4.1–4.3 → style-profile.json
    profile = {
        "page": extract_page(doc),
        "style_roles": roles,
        "table_styles": extract_table_styles(doc, normal_size),
        "title_block": extract_title_block(doc, normal_size),
        "block_patterns": extract_block_patterns(doc, texts),
        "microtypography": scan_microtypography(texts),
        "table_templates": extract_table_templates(doc),
        "section_blueprint": extract_section_blueprint(doc),
        "glossary_terms": build_glossary(texts),
    }
    errs = validate_profile(profile)
    if errs:
        print("[stop] профиль не валиден против схемы:", file=sys.stderr)
        for e in errs:
            print("  -", e, file=sys.stderr)
        return 1
    profile_path = outdir / "style-profile.json"
    profile_path.write_text(json.dumps(profile, ensure_ascii=False, indent=2) + "\n",
                            encoding="utf-8")

    # Step 4.4 → template.docx
    template_path = build_template_docx(sample, outdir)

    # Step 4.5 → черновики
    examples = profile["microtypography"].get("examples", {})
    thesis_example = next((t for t in texts if not is_lead_in_pseudo(t)
                           and len(t.split()) >= 6), texts[0] if texts else "—")
    lead = profile["block_patterns"]["lead_in_count"]
    caption_re = profile["block_patterns"]["figure_caption_regex"]
    lead_example = next((t for t in texts if t.rstrip().endswith(":")), None)
    formula_example = next((t for p in doc.paragraphs
                            if p.style.name == "Normal" and is_fully_italic(p)
                            for t in [p.text]), None)
    bullet_example = next((p.text for p in doc.paragraphs if p.style.name == "List Bullet"), None)
    table_example = " | ".join(profile["table_templates"][0]["header"]) \
        if profile["table_templates"] else None
    caption_example = next((t for t in texts if re.search(caption_re, t)), None)
    verdict_example = None
    for t in texts:
        m = re.search(r".{0,60}\b(ПОДТВЕРЖДЕНА|ОТВЕРГНУТА|НЕ ПОДТВЕРЖДЕНА|ПРИНЯТА)\b.{0,60}", t)
        if m:
            verdict_example = m.group(0)
            break
    block_examples = {
        "thesis": thesis_example,
        "lead_in": lead_example,
        "formula": formula_example,
        "bullets": bullet_example,
        "table": table_example,
        "figure": caption_example,
        "verdict": verdict_example,
    }

    # first_person — честный прогон категории по текстам образца
    first_person = 0
    if base_blacklist.is_file():
        for line in base_blacklist.read_text(encoding="utf-8").splitlines():
            if line.strip() and not line.startswith("#") and line.startswith("first_person\t"):
                rx = re.compile(line.split("\t", 1)[1].strip(), re.IGNORECASE)
                first_person = sum(len(rx.findall(t)) for t in texts)
                break

    write_styleguide_draft(outdir / "styleguide.draft.md", profile, sample, texts,
                           examples, block_examples, first_person)
    write_glossary_draft(outdir / "glossary.draft.md", sample,
                         profile["glossary_terms"], texts)

    # Step 4.6 → work/slang-blacklist.txt
    legal_commented = []
    if base_blacklist.is_file():
        legal_commented = copy_blacklist(base_blacklist, outdir / "slang-blacklist.txt",
                                         profile["glossary_terms"])
    else:
        print(f"[warn] база блэклиста не найдена: {base_blacklist}", file=sys.stderr)

    # Краткая сводка в stdout
    micro = profile["microtypography"]
    print("=" * 72)
    print(f"Извлечение профиля: {sample.name}")
    print("=" * 72)
    pg = profile["page"]
    print(f"Страница: {pg['width_emu']}x{pg['height_emu']} EMU, {pg['orientation']}, "
          f"поля {pg['margins_emu']['left']}/{pg['margins_emu']['top']} EMU (L/T)")
    print("Роли стилей:")
    for key, r in profile["style_roles"].items():
        print(f"  {key:12s} font={r['font']!r} (source={r['source']}) size={r['size_pt']} "
              f"bold={r['bold']} color={r['color_hex']} before={r['spacing_before_emu']} "
              f"after={r['spacing_after_emu']}")
    print(f"Стили таблиц: {[ts['name'] for ts in profile['table_styles']]} "
          f"(cell={profile['table_styles'][0]['cell_size_pt']}pt, "
          f"header_bold={profile['table_styles'][0]['header_bold']})")
    print(f"Титульный блок: {len(profile['title_block'])} абзацев: "
          f"{[(t['size_pt'], t['bold'], t['italic'], t['alignment']) for t in profile['title_block']]}")
    print(f"Блоковые паттерны: caption={profile['block_patterns']['figure_caption_regex']!r} "
          f"(матчей {profile['block_patterns']['figure_caption_count']}), "
          f"formula={profile['block_patterns']['formula_style']} "
          f"(курсивных абзацев {profile['block_patterns']['formula_count']}), "
          f"lead_in={profile['block_patterns']['lead_in_terminator']!r} "
          f"(вхождений {profile['block_patterns']['lead_in_count']})")
    print("Микротипографика (детерминанты):")
    for key in ("decimal_separator", "minus_char", "percent_spacing", "range_dash", "quotes"):
        print(f"  {key}: {micro[key]}")
    print(f"  arrow={micro['arrow']} approx={micro['approx']} multiplication={micro['multiplication']}")
    print(f"  счётчики: {dict(micro['patterns'])}")
    print(f"Шаблоны таблиц: {len(profile['table_templates'])} "
          f"({', '.join(t['name'] + ' x' + str(t['occurrences']) for t in profile['table_templates'])})")
    h1 = [s["heading"] for s in profile["section_blueprint"] if s["level"] == 1]
    print(f"Blueprint разделов: всего {len(profile['section_blueprint'])}, H1: {len(h1)}")
    for h in h1:
        print(f"  - {h}")
    print(f"Глоссарий: {len(profile['glossary_terms'])} терминов; топ-10: "
          f"{[(g['term'], g['count']) for g in profile['glossary_terms'][:10]]}")
    print(f"1-го лица в текстах образца: {first_person}")
    print(f"Артефакты в {outdir}:")
    for f in sorted(outdir.iterdir()):
        print(f"  {f.name} ({f.stat().st_size} байт)")
    if legal_commented:
        print(f"Блэклист: закомментированы категории: {legal_commented}")
    else:
        print("Блэклист: пересечений глоссария с категориями нет — копия без изменений")
    return 0


def is_lead_in_pseudo(text: str) -> bool:
    """Текст сам по себе — lead-in (жирность недоступна по тексту): используется
    только для выбора примера тезиса (исключаем строки-подводки с «:»)."""
    return text.rstrip().endswith(":")


if __name__ == "__main__":
    sys.exit(main())
