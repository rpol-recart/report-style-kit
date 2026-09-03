#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_report.py — Task 5 (Steps 5.1–5.2) комплекта report-style-kit.

Собирает .docx-отчёт из content-manifest на базе template.docx и style-profile.json.
Стиль гарантируется наследованием от шаблона: скрипт не изобретает оформление,
а детерминированно маппит блоковую грамматику манифеста на стили, извлечённые
из эталонного отчёта.

CLI:
    python3 build_report.py --manifest m.json --profile p.json --template t.docx -o out.docx

Коды выхода:
    0 — отчёт собран;
    2 — манифест не прошёл валидацию (ошибки по-русски, с путями вида
        sections[0].blocks[1].text);
    3 — не найден файл манифеста/профиля/шаблона;
    4 — отчёт собран, но часть картинок не найдена (список в stdout,
        на их месте вставлен пустой абзац).

Зависимости: python-docx (остальное — стандартная библиотека). Python 3.12.
"""
from __future__ import annotations

import argparse
import json
import sys
from collections import Counter
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Emu, Pt

BLOCK_TYPES = ("thesis", "lead_in", "formula", "bullets", "table", "figure", "verdict")
META_FIELDS = ("title", "subtitle", "project", "version", "period")

# Ключи блоков по типам — зеркало additionalProperties из content-manifest.schema.json.
BLOCK_KEYS = {
    "thesis": {"type", "text"},
    "lead_in": {"type", "label", "text"},
    "formula": {"type", "text"},
    "bullets": {"type", "items"},
    "table": {"type", "header", "rows", "template"},
    "figure": {"type", "caption", "image", "explanation"},
    "verdict": {"type", "text"},
}

# Запасной титульный паттерн, если в профиле нет title_block (не рекомендуемый путь).
DEFAULT_TITLE_BLOCK = [
    {"size_pt": 20.0, "bold": True, "italic": False, "alignment": "center"},
    {"size_pt": 14.0, "bold": False, "italic": True, "alignment": "center"},
    {"size_pt": 11.0, "bold": False, "italic": False, "alignment": "center"},
]


# --------------------------------------------------------------------------- validation (Step 5.1)

def _is_str(v) -> bool:
    return isinstance(v, str) and v.strip() != ""


def _check_str(errs: list, path: str, value, what: str = "непустая строка") -> None:
    if not _is_str(value):
        errs.append(f"{path}: отсутствует или не {what} (получено {value!r})")


def _table_template_header(profile: dict, name) -> list | None:
    for t in profile.get("table_templates") or []:
        if t.get("name") == name:
            header = t.get("header")
            return header if isinstance(header, list) else None
    return None


def _validate_table_block(b: dict, bp: str, errs: list, profile: dict) -> list | None:
    """Валидация блока table; возвращает разрешённый header (или None при ошибке)."""
    template = b.get("template")
    header = b.get("header")
    rows = b.get("rows")
    resolved: list | None = header if isinstance(header, list) else None

    if template is not None:
        if not _is_str(template):
            errs.append(f"{bp}.template: должно быть непустой строкой (именем шаблона)")
        else:
            th = _table_template_header(profile, template)
            if th is None:
                avail = ", ".join(sorted(
                    t.get("name", "") for t in (profile.get("table_templates") or [])
                )) or "— нет зарегистрированных шаблонов —"
                errs.append(f"{bp}.template: шаблон {template!r} не найден в "
                            f"profile.table_templates (доступные шаблоны: {avail})")
            elif resolved is None:
                resolved = th
    elif not isinstance(header, list) or not header or not all(_is_str(h) for h in header):
        errs.append(f"{bp}.header: без 'template' обязателен непустой массив строк")

    if resolved is not None:
        if "template" not in b and (not isinstance(header, list) or not header):
            resolved = None
    else:
        # header неизвестен: проверить только форму rows
        if rows is not None and (not isinstance(rows, list)
                                 or not all(isinstance(r, list) for r in rows)):
            errs.append(f"{bp}.rows: должен быть массивом массивов строк")

    if "template" not in b and (not isinstance(rows, list) or not rows):
        errs.append(f"{bp}.rows: без 'template' обязателен непустой массив рядов "
                    f"(с 'template' ряды могут отсутствовать)")
    elif isinstance(rows, list) and resolved is not None:
        for r, row in enumerate(rows):
            if not isinstance(row, list) or not all(isinstance(c, str) for c in row):
                errs.append(f"{bp}.rows[{r}]: должен быть массивом строк")
            elif len(row) != len(resolved):
                errs.append(f"{bp}.rows[{r}]: длина ряда ({len(row)}) не совпадает "
                            f"с числом колонок ({len(resolved)})")
    return resolved if isinstance(resolved, list) else None


def validate_manifest(m, profile: dict) -> list:
    """Ручной валидатор (без внешней jsonschema). Возвращает список ошибок по-русски
    с путями вида sections[3].blocks[2].text."""
    if not isinstance(m, dict):
        return ["манифест: должен быть JSON-объектом"]
    errs: list = []

    meta = m.get("meta")
    if not isinstance(meta, dict):
        errs.append("meta: отсутствует обязательный раздел meta (объект с title, subtitle, "
                    "project, version, period)")
    else:
        for f in META_FIELDS:
            _check_str(errs, f"meta.{f}", meta.get(f))

    sections = m.get("sections")
    if not isinstance(sections, list) or not sections:
        errs.append("sections: отсутствует или не непустой массив разделов")
        return errs

    for i, sec in enumerate(sections):
        sp = f"sections[{i}]"
        if not isinstance(sec, dict):
            errs.append(f"{sp}: раздел должен быть объектом {{heading, level, blocks}}")
            continue
        _check_str(errs, f"{sp}.heading", sec.get("heading"))
        level = sec.get("level")
        if level not in (1, 2, 3):
            errs.append(f"{sp}.level: должно быть 1, 2 или 3 (получено {level!r})")
        blocks = sec.get("blocks")
        if not isinstance(blocks, list):
            errs.append(f"{sp}.blocks: отсутствует массив блоков")
            continue
        for j, b in enumerate(blocks):
            bp = f"{sp}.blocks[{j}]"
            if not isinstance(b, dict):
                errs.append(f"{bp}: блок должен быть объектом с полем 'type' "
                            f"({', '.join(BLOCK_TYPES)})")
                continue
            btype = b.get("type")
            if btype not in BLOCK_KEYS:
                errs.append(f"{bp}.type: неизвестный тип блока {btype!r} "
                            f"(допустимо: {', '.join(BLOCK_TYPES)})")
                continue
            extra = set(b) - BLOCK_KEYS[btype]
            if extra:
                errs.append(f"{bp}: недопустимые поля для типа {btype!r}: "
                            f"{', '.join(sorted(extra))}")
            if btype in ("thesis", "formula", "verdict"):
                _check_str(errs, f"{bp}.text", b.get("text"))
            elif btype == "lead_in":
                _check_str(errs, f"{bp}.label", b.get("label"))
                _check_str(errs, f"{bp}.text", b.get("text"))
            elif btype == "bullets":
                items = b.get("items")
                if not isinstance(items, list) or not items \
                        or not all(_is_str(it) for it in items):
                    errs.append(f"{bp}.items: обязателен непустой массив непустых строк")
            elif btype == "table":
                _validate_table_block(b, bp, errs, profile)
            elif btype == "figure":
                _check_str(errs, f"{bp}.caption", b.get("caption"))
                _check_str(errs, f"{bp}.image", b.get("image"))
                if "explanation" in b:
                    _check_str(errs, f"{bp}.explanation", b.get("explanation"))
    return errs


# --------------------------------------------------------------------------- build (Step 5.2)

def _fill_cell(cell, text: str, font_name: str, size_pt: float, bold) -> None:
    run = cell.paragraphs[0].add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bool(bold)


def build(manifest: dict, profile: dict, template_path: Path, out_path: Path,
          base_dir: Path | None = None) -> dict:
    doc = docx.Document(str(template_path))
    meta = manifest["meta"]
    base_dir = base_dir or Path.cwd()  # относительные пути картинок — от каталога манифеста

    # --- титульный блок по profile.title_block -------------------------------
    title_block = profile.get("title_block") or DEFAULT_TITLE_BLOCK
    title_texts = [
        meta["title"],
        meta["subtitle"],
        f"Проект: {meta['project']}  |  Версия: {meta['version']}\n"
        f"Период данных: {meta['period']}",
    ]
    for i, text in enumerate(title_texts):
        el = title_block[i] if i < len(title_block) else title_block[-1]
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(el.get("size_pt") or 11.0)
        run.bold = el.get("bold")
        run.italic = el.get("italic")
    doc.add_paragraph("", style="Normal")  # один пустой абзац после титула

    # --- параметры стиля из профиля ------------------------------------------
    roles = profile.get("style_roles") or {}
    normal_font = (roles.get("normal") or {}).get("font") or "Calibri"
    tstyle = (profile.get("table_styles") or [{}])[0]
    cell_size = tstyle.get("cell_size_pt") or 10.0
    header_bold = tstyle.get("header_bold")
    header_bold = True if header_bold is None else bool(header_bold)
    page = profile.get("page") or {}
    margins = page.get("margins_emu") or {}
    avail_emu = int(page.get("width_emu") or 0) \
        - int(margins.get("left") or 0) - int(margins.get("right") or 0)
    templates = {t.get("name"): t.get("header") or []
                 for t in profile.get("table_templates") or []}

    counts: Counter = Counter()
    n_tables = 0
    n_pictures = 0
    missing_images: list = []

    for i, sec in enumerate(manifest["sections"]):
        doc.add_paragraph(sec["heading"], style=f"Heading {sec['level']}")
        for j, b in enumerate(sec["blocks"]):
            btype = b["type"]
            counts[btype] += 1
            bp = f"sections[{i}].blocks[{j}]"

            if btype == "thesis":
                doc.add_paragraph(b["text"], style="Normal")
            elif btype == "verdict":
                doc.add_paragraph(b["text"], style="List Bullet")
            elif btype == "lead_in":
                # Паттерн эталона: жирный label — отдельный абзац, текст —
                # следующим абзацем (как «Бизнес-вопрос:» + вопрос в образце).
                plabel = doc.add_paragraph(style="Normal")
                plabel.add_run(b["label"]).bold = True
                if b["text"]:
                    doc.add_paragraph(b["text"], style="Normal")
            elif btype == "formula":
                p = doc.add_paragraph(style="Normal")
                p.add_run(b["text"]).italic = True
            elif btype == "bullets":
                for item in b["items"]:
                    doc.add_paragraph(item, style="List Bullet")
            elif btype == "table":
                header = b.get("header") or templates.get(b.get("template")) or []
                rows = b.get("rows") or []
                table = doc.add_table(rows=1 + len(rows), cols=len(header))
                table.style = tstyle.get("name") or "Table Grid"
                table.autofit = True
                for c, htxt in enumerate(header):
                    _fill_cell(table.cell(0, c), htxt, normal_font, cell_size, header_bold)
                for r, row in enumerate(rows, start=1):
                    for c, ctxt in enumerate(row):
                        _fill_cell(table.cell(r, c), ctxt, normal_font, cell_size, False)
                n_tables += 1
            elif btype == "figure":
                doc.add_paragraph(b["caption"], style="Normal")
                ppic = doc.add_paragraph(style="Normal")
                img = Path(b["image"])
                if not img.is_absolute():
                    img = base_dir / img
                if img.is_file():
                    run = ppic.add_run()
                    run.add_picture(str(img),
                                    width=Emu(avail_emu) if avail_emu > 0 else None)
                    ppic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    n_pictures += 1
                else:
                    # пустой абзац вместо картинки; нарушение фиксируется в отчёте
                    missing_images.append((bp, b["image"]))
                if b.get("explanation"):
                    doc.add_paragraph(b["explanation"], style="Normal")

    doc.save(str(out_path))
    return {
        "counts": counts,
        "n_tables": n_tables,
        "n_pictures": n_pictures,
        "missing_images": missing_images,
    }


# --------------------------------------------------------------------------- main

def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(
        description="Сборка docx-отчёта из content-manifest по стилевому профилю.")
    ap.add_argument("--manifest", required=True, help="content-manifest.json")
    ap.add_argument("--profile", required=True, help="style-profile.json (из extract_profile.py)")
    ap.add_argument("--template", required=True, help="template.docx (из extract_profile.py)")
    ap.add_argument("-o", "--out", required=True, help="путь итогового .docx")
    args = ap.parse_args(argv)

    for path, what in ((args.manifest, "манифест"), (args.profile, "профиль стиля"),
                       (args.template, "шаблон")):
        if not Path(path).is_file():
            print(f"[ошибка] {what} не найден: {path}", file=sys.stderr)
            return 3

    try:
        manifest = json.loads(Path(args.manifest).read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        print(f"[ошибка] манифест — некорректный JSON: {exc}", file=sys.stderr)
        return 2
    try:
        profile = json.loads(Path(args.profile).read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        print(f"[ошибка] профиль — некорректный JSON: {exc}", file=sys.stderr)
        return 3

    if not isinstance(profile, dict):
        print("[ошибка] профиль должен быть JSON-объектом", file=sys.stderr)
        return 3
    errs = validate_manifest(manifest, profile)
    if errs:
        print("[ошибка] манифест не прошёл валидацию:", file=sys.stderr)
        for e in errs:
            print(f"  - {e}", file=sys.stderr)
        return 2

    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    result = build(manifest, profile, Path(args.template), out_path,
                   base_dir=Path(args.manifest).resolve().parent)

    total_blocks = sum(result["counts"].values())
    by_type = ", ".join(f"{t}: {n}" for t, n in result["counts"].items()) or "—"
    print("=" * 72)
    print(f"Отчёт собран: {out_path}")
    print(f"Секций: {len(manifest['sections'])}")
    print(f"Блоков: {total_blocks} ({by_type})")
    print(f"Таблиц: {result['n_tables']}")
    print(f"Картинок: {result['n_pictures']} "
          f"(пропущено: {len(result['missing_images'])})")
    print("=" * 72)
    if result["missing_images"]:
        print("[ошибка] файлы картинок не найдены (в отчёт вставлен пустой абзац):",
              file=sys.stderr)
        for loc, img in result["missing_images"]:
            print(f"  - {loc}.image: {img}", file=sys.stderr)
        return 4
    return 0


if __name__ == "__main__":
    sys.exit(main())
