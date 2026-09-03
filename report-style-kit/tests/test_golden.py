#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
test_golden.py — Task 9.1 комплекта report-style-kit. Golden-клон эталона.

Шаги:
  1. extract_profile.py на эталоне → tests/work_golden/ (subprocess, cwd — корень
     комплекта, пути аргументов относительные).
  2. Программное построение golden_manifest.json из самого эталона (python-docx,
     обход body в порядке документа):
       - титульные абзацы до первого Heading          → meta (title/subtitle,
         project/version/period разбираются из третьей строки);
       - Heading 1/2/3                                 → sections level 1/2/3
         (текст заголовка как есть);
       - подряд идущие List Bullet                     → один блок bullets{items};
       - полностью курсивный непустой Normal           → formula{text};
       - абзац, начинающийся с bold-run на «:»         → lead_in{label, text},
         где text — следующий за ним простой Normal-абзац (поглощается); если
         дальше идёт не простой абзац (таблица/буллеты/формула) — блок thesis
         с текстом жирной подводки;
       - Normal c «Рис. N. …»                          → caption фигуры;
       - абзац с a:blip                                → figure{caption, image};
         image — извлечённая i-я картинка media эталона (image{i}.png →
         tests/work_golden/), i — порядковый номер figure-блока; следующий
         простой Normal-абзац — explanation;
       - таблицы                                       → table{header=ряд 0,
         rows=остальные} на своей позиции в body;
       - прочие непустые Normal                        → thesis{text}.
     Абсолютные пути хоста, встретившиеся в текстах эталона, маскируются в
     «<путь-образца>» (аудит copy-out; на стили не влияет).
  3. build_report.py (профиль/шаблон из tests/work_golden/) → golden_built.docx.
  4. lint_visual.py на golden_built.docx с профилем → обязан быть exit 0.
  5. Стилевой diff golden_built.docx ↔ эталон: последовательности непустых
     абзацев (style.name) и таблиц (style.name + размерность) в порядке body,
     попарно; титульный блок — первые 3 непустых абзаца по (size/bold/italic/
     alignment первого run). Пустые абзацы пропускаются.

Итог: 0 расхождений → «GOLDEN PASS: 0 style diffs», exit 0;
иначе «GOLDEN FAIL: N diffs» + список (≤20), exit 1.

Запуск: python3 tests/test_golden.py   (из любого каталога).
Зависимости: python-docx (эталон только читается).
"""
from __future__ import annotations

import json
import re
import subprocess
import sys
import zipfile
from pathlib import Path

import docx
from docx.oxml.ns import qn

KIT = Path(__file__).resolve().parents[1]          # корень комплекта
SAMPLE = KIT.parent / "CR_Predictor_Report.docx"   # эталон (только чтение)
WORK = KIT / "tests" / "work_golden"
WORK_REL = "tests/work_golden"                     # относительный путь от KIT

HEADING_STYLES = {"Heading 1": 1, "Heading 2": 2, "Heading 3": 3}
CAPTION_RE = re.compile(r"^Рис\. \d+\.")
# Абсолютные пути хоста внутри текстов эталона (та же маскировка, что в
# extract_profile.sanitize_sample_text; продублировано, чтобы тест был
# автономным исполняемым скриптом).
ABS_PATH_RE = re.compile(r"(?<![\w])/(?:root|home|Users)/[^\s|»«()]*")

MAX_DIFFS_SHOWN = 20


def sanitize(text: str) -> str:
    return ABS_PATH_RE.sub("<путь-образца>", text)


# --------------------------------------------------------------------- шаг 1

def run_step(cmd: list, what: str) -> None:
    proc = subprocess.run([sys.executable] + cmd, cwd=KIT,
                          capture_output=True, text=True)
    if proc.returncode != 0:
        print(f"[golden] шаг не прошёл: {what} (exit {proc.returncode})",
              file=sys.stderr)
        print("--- stdout ---\n" + proc.stdout, file=sys.stderr)
        print("--- stderr ---\n" + proc.stderr, file=sys.stderr)
        raise SystemExit(1)
    return proc


# --------------------------------------------------------------------- шаг 2

def has_blip(p) -> bool:
    return bool(p._p.findall(".//" + qn("a:blip")))


def is_fully_italic(p) -> bool:
    runs = [r for r in p.runs if r.text.strip()]
    return bool(runs) and all(r.italic is True for r in runs)


def is_bold_lead(p) -> bool:
    return bool(p.runs) and bool(p.runs[0].bold) \
        and p.runs[0].text.rstrip().endswith(":")


def is_plain_normal(p) -> bool:
    """Простой текстовый Normal-абзац (кандидат в текст lead_in/explanation)."""
    return p.style.name == "Normal" and bool(p.text.strip()) \
        and not has_blip(p) and not is_bold_lead(p) and not is_fully_italic(p) \
        and not CAPTION_RE.match(p.text.strip())


def body_items(doc):
    """[(kind, obj)] всех детей body в порядке документа."""
    para_map = {p._p: p for p in doc.paragraphs}
    table_map = {t._tbl: t for t in doc.tables}
    items = []
    for child in doc.element.body:
        if child.tag == qn("w:p") and child in para_map:
            items.append(("p", para_map[child]))
        elif child.tag == qn("w:tbl") and child in table_map:
            items.append(("t", table_map[child]))
    return items


def extract_meta(items) -> dict:
    """Титульные непустые абзацы до первого Heading → meta."""
    title_paras = []
    for kind, obj in items:
        if kind == "p":
            if obj.style.name in HEADING_STYLES:
                break
            if obj.text.strip():
                title_paras.append(obj.text)
        else:
            break
    if len(title_paras) < 3:
        raise SystemExit(f"[golden] в эталоне меньше 3 титульных абзацев: "
                         f"{len(title_paras)}")
    line = title_paras[2]
    m_project = re.search(r"Проект:\s*(.+?)(?:\s*\||\n|$)", line)
    m_version = re.search(r"Версия(?:\s+\w+)?\s*:\s*(.+?)(?:\n|$)", line)
    m_period = re.search(r"Период данных:\s*(.+?)(?:\s*\||\n|$)", line)
    if not (m_project and m_version and m_period):
        raise SystemExit(f"[golden] не удалось разобрать строку титула: {line!r}")
    return {
        "title": title_paras[0].strip(),
        "subtitle": title_paras[1].strip(),
        "project": m_project.group(1).strip(),
        "version": m_version.group(1).strip(),
        "period": m_period.group(1).strip(),
    }


def extract_media(sample: Path, figure_count: int) -> None:
    """word/media/image{i}.png эталона → tests/work_golden/image{i}.png."""
    with zipfile.ZipFile(sample) as z:
        media = {}
        for name in z.namelist():
            m = re.fullmatch(r"word/media/image(\d+)\.(?:png|jpeg|jpg)", name)
            if m:
                media[int(m.group(1))] = name
        ordered = [media[k] for k in sorted(media)]
        if len(ordered) < figure_count:
            raise SystemExit(f"[golden] картинок в media ({len(ordered)}) меньше, "
                             f"чем figure-блоков ({figure_count})")
        for i in range(1, figure_count + 1):
            (WORK / f"image{i}.png").write_bytes(z.read(ordered[i - 1]))


def build_manifest(items, sample: Path) -> dict:
    meta = extract_meta(items)
    sections: list = []
    cur: dict | None = None
    pending_caption: str | None = None
    expect_explanation = False
    figure_count = 0

    def push(block: dict) -> None:
        blocks = cur["blocks"]
        if block["type"] == "bullets" and blocks and blocks[-1]["type"] == "bullets":
            blocks[-1]["items"].extend(block["items"])
        else:
            blocks.append(block)

    idx = 0
    seen_heading = False
    while idx < len(items):
        kind, obj = items[idx]
        idx += 1
        if kind == "t":
            if not seen_heading:
                continue  # таблицы в титульной зоне не ожидаются
            rows = [[c.text for c in row.cells] for row in obj.rows]
            push({"type": "table", "header": rows[0], "rows": rows[1:]})
            continue

        p = obj
        sname = p.style.name
        if sname in HEADING_STYLES:
            seen_heading = True
            cur = {"heading": p.text.strip(), "level": HEADING_STYLES[sname],
                   "blocks": []}
            sections.append(cur)
            pending_caption = None
            expect_explanation = False
            continue
        if not seen_heading:
            continue  # титульная зона разобрана extract_meta
        if has_blip(p):
            figure_count += 1
            block = {"type": "figure", "caption": pending_caption or "",
                     "image": f"image{figure_count}.png"}
            pending_caption = None
            expect_explanation = True
            push(block)
            continue
        if not p.text.strip():
            continue  # пустые абзацы-разделители
        if CAPTION_RE.match(p.text.strip()):
            pending_caption = sanitize(p.text.strip())
            continue
        if sname == "List Bullet":
            push({"type": "bullets", "items": [sanitize(p.text.strip())]})
            continue
        if is_bold_lead(p):
            nxt = items[idx] if idx < len(items) else None
            if nxt and nxt[0] == "p" and is_plain_normal(nxt[1]):
                push({"type": "lead_in", "label": sanitize(p.text.strip()),
                      "text": sanitize(nxt[1].text.strip())})
                idx += 1  # текст подводки поглощён
                continue
            push({"type": "thesis", "text": sanitize(p.text.strip())})
            continue
        # простой Normal: сначала объяснение к предыдущей фигуре (оно может быть
        # полностью курсивным — в эталоне пояснения под рисунками курсивом),
        # затем формула
        if expect_explanation and cur and cur["blocks"] \
                and cur["blocks"][-1]["type"] == "figure" \
                and "explanation" not in cur["blocks"][-1]:
            cur["blocks"][-1]["explanation"] = sanitize(p.text.strip())
            expect_explanation = False
            continue
        if is_fully_italic(p):
            push({"type": "formula", "text": sanitize(p.text.strip())})
            continue
        push({"type": "thesis", "text": sanitize(p.text.strip())})

    if not sections:
        raise SystemExit("[golden] в эталоне не найдено ни одного Heading-раздела")
    extract_media(sample, figure_count)
    return {"meta": meta, "sections": sections}


# --------------------------------------------------------------------- шаг 5

def style_sequence(doc):
    """[(('p', style.name) | ('t', style.name, rows, cols))] непустых абзацев
    и всех таблиц в порядке body."""
    para_map = {p._p: p for p in doc.paragraphs}
    table_map = {t._tbl: t for t in doc.tables}
    seq = []
    for child in doc.element.body:
        if child.tag == qn("w:p") and child in para_map:
            p = para_map[child]
            if p.text.strip():
                seq.append(("p", p.style.name))
        elif child.tag == qn("w:tbl") and child in table_map:
            t = table_map[child]
            try:
                sname = t.style.name if t.style is not None else None
            except KeyError:
                sname = None
            seq.append(("t", sname, len(t.rows), len(t.columns)))
    return seq


def title_signature(doc, count: int = 3):
    """Первые `count` непустых абзацев: (size, bold, italic, alignment) 1-го run."""
    out = []
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        if len(out) >= count:
            break
        run = p.runs[0] if p.runs else None
        out.append((
            run.font.size.pt if run is not None and run.font.size is not None else None,
            run.bold if run is not None else None,
            run.italic if run is not None else None,
            p.alignment.name if p.alignment is not None else None,
        ))
    return out


# ----------------------------------------------------------------------- main

def main() -> int:
    if not SAMPLE.is_file():
        print(f"[golden] эталон не найден: {SAMPLE}", file=sys.stderr)
        return 1
    WORK.mkdir(parents=True, exist_ok=True)

    # --- шаг 1: профиль + шаблон из эталона ----------------------------------
    run_step(["scripts/extract_profile.py",
              "--sample", "../CR_Predictor_Report.docx",
              "--outdir", WORK_REL],
             "extract_profile.py")

    # --- шаг 2: golden-манифест из эталона -----------------------------------
    ref_doc = docx.Document(str(SAMPLE))
    manifest = build_manifest(body_items(ref_doc), SAMPLE)
    manifest_path = WORK / "golden_manifest.json"
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n",
                             encoding="utf-8")

    # --- шаг 3: сборка ---------------------------------------------------------
    run_step(["scripts/build_report.py",
              "--manifest", f"{WORK_REL}/golden_manifest.json",
              "--profile", f"{WORK_REL}/style-profile.json",
              "--template", f"{WORK_REL}/template.docx",
              "-o", f"{WORK_REL}/golden_built.docx"],
             "build_report.py")

    # --- шаг 4: визуальный линт построенного ---------------------------------
    proc = run_step(["scripts/lint_visual.py",
                     "--docx", f"{WORK_REL}/golden_built.docx",
                     "--profile", f"{WORK_REL}/style-profile.json"],
                    "lint_visual.py")

    # --- шаг 5: стилевой diff --------------------------------------------------
    built_doc = docx.Document(str(WORK / "golden_built.docx"))
    diffs: list = []

    ref_seq, built_seq = style_sequence(ref_doc), style_sequence(built_doc)
    for i, (r, b) in enumerate(zip(ref_seq, built_seq)):
        if r != b:
            diffs.append(f"элемент {i}: эталон {r} vs построен {b}")
    if len(ref_seq) != len(built_seq):
        diffs.append(f"число элементов: эталон {len(ref_seq)} vs построен {len(built_seq)}")

    ref_title, built_title = title_signature(ref_doc), title_signature(built_doc)
    for k, (r, b) in enumerate(zip(ref_title, built_title)):
        if r != b:
            diffs.append(f"титул абзац {k} (size/bold/italic/align): "
                         f"эталон {r} vs построен {b}")

    if diffs:
        print(f"GOLDEN FAIL: {len(diffs)} diffs")
        for d in diffs[:MAX_DIFFS_SHOWN]:
            print("  -", d)
        if len(diffs) > MAX_DIFFS_SHOWN:
            print(f"  … и ещё {len(diffs) - MAX_DIFFS_SHOWN}")
        return 1

    print("GOLDEN PASS: 0 style diffs")
    print(f"lint_visual: {proc.stdout.strip().splitlines()[-1]}")
    print(f"Абзацев: эталон {len(ref_seq)}, построен {len(built_seq)}; "
          f"секций: {len(manifest['sections'])}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
