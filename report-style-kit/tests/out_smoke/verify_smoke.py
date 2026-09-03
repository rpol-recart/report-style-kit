#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Смоук-проверка smoke.docx (Task 5) через python-docx."""
import sys
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

KIT = Path(__file__).resolve().parents[2]
DOC = KIT / "tests" / "out_smoke" / "smoke.docx"
PROFILE = KIT / "tests" / "work_extract_check" / "style-profile.json"

import json
profile = json.loads(PROFILE.read_text(encoding="utf-8"))
avail = profile["page"]["width_emu"] - profile["page"]["margins_emu"]["left"] \
    - profile["page"]["margins_emu"]["right"]

d = docx.Document(str(DOC))
failures = []


def check(cond, msg):
    print(("OK  " if cond else "FAIL") + " " + msg)
    if not cond:
        failures.append(msg)


paras = d.paragraphs
styles = [p.style.name for p in paras]
print("Стили абзацев:", styles)
texts = [p.text for p in paras]

# Титульный блок: 3 абзаца Normal по центру + пустой Normal
check(styles[:3] == ["Normal"] * 3, "титульные 3 абзаца — Normal")
check(texts[0] == "Отчёт: синтетический смоук", "title = meta.title")
check(texts[1] == "Проверка конвейера сборки", "subtitle = meta.subtitle")
check(texts[2] == "Проект: report-style-kit  |  Версия: 1.0\nПериод данных: 2026-09"
      .replace("\n", "\n"), f"строка проекта/версии/периода: {texts[2]!r}")
check(styles[3] == "Normal" and texts[3] == "", "после титула один пустой Normal")
for i, want in enumerate([(20.0, True, None), (14.0, None, True), (11.0, None, None)]):
    r = paras[i].runs[0]
    size, bold, italic = want
    check(r.font.size.pt == size and r.bold is bold and r.italic is italic,
          f"title_block[{i}]: size={r.font.size.pt} bold={r.bold} italic={r.italic}")
    check(paras[i].alignment == WD_ALIGN_PARAGRAPH.CENTER, f"title_block[{i}]: alignment=CENTER")

# Секции
check(styles[4] == "Heading 1" and texts[4] == "1. Методика", "заголовок 1 секции")
check(styles[5] == "Normal" and not paras[5].runs[0].bold, "thesis → Normal")
check(styles[6] == "Normal" and paras[6].runs[0].bold is True
      and paras[6].runs[0].text == "Входные данные: "
      and paras[6].runs[1].bold is not True, "lead_in → Normal, label bold, text без bold")
check(paras[6].runs[1].italic is not True, "lead_in text не курсив")
check(styles[7] == "Normal" and all(r.italic for r in paras[7].runs),
      "formula → Normal, весь run italic")
check(styles[8] == "List Bullet" and styles[9] == "List Bullet"
      and styles[10] == "List Bullet", "bullets ×3 → List Bullet")
check(styles[11] == "Normal" and texts[11] == "Рис. 1. Тестовая гистограмма.",
      "figure: caption → Normal")
check(styles[13] == "Normal"
      and texts[13] == "На графике видно снижение ошибки после перехода на широкую сетку.",
      "figure: explanation → Normal")
check(styles[14] == "Heading 1" and texts[14] == "2. Выводы", "заголовок 2 секции")
check(styles[15] == "List Bullet", "verdict → List Bullet")

# Картинка
check(len(d.inline_shapes) == 1, "inline_shapes == 1")
shp = d.inline_shapes[0]
check(int(shp.width) == avail, f"ширина картинки {int(shp.width)} EMU == доступная {avail}")
pic_para = paras[12]
check(pic_para.alignment == WD_ALIGN_PARAGRAPH.CENTER, "абзац с картинкой по центру")

# Таблица
check(len(d.tables) == 1, "таблиц == 1")
t = d.tables[0]
check(t.style.name == "Light Grid Accent 1", f"стиль таблицы = {t.style.name!r}")
check(len(t.rows) == 3 and len(t.columns) == 3, f"размерность 3x3 ({len(t.rows)}x{len(t.columns)})")
hdr = [c.text for c in t.rows[0].cells]
check(hdr == ["Метод", "Avg MAE", "Δ vs baseline"], f"header: {hdr}")
for c in t.rows[0].cells:
    r = c.paragraphs[0].runs[0]
    check(r.bold is True, f"header ячейка {c.text!r}: bold")
    check(r.font.size.pt == 10.0, f"header ячейка {c.text!r}: {r.font.size.pt}pt")
    check(r.font.name == "Calibri", f"header ячейка {c.text!r}: font={r.font.name}")
for ri in (1, 2):
    for c in t.rows[ri].cells:
        r = c.paragraphs[0].runs[0]
        check(r.font.size.pt == 10.0 and r.font.name == "Calibri" and r.bold is not True,
              f"data[{ri}] {c.text!r}: 10pt Calibri, не bold")
check(t.autofit, "table.autofit")

print()
if failures:
    print(f"VERIFY FAIL: {len(failures)} проверок провалено")
    sys.exit(1)
print("VERIFY PASS: все проверки пройдены")
